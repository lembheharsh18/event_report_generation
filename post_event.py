import streamlit as st
from groq import Groq
from docx import Document
import io
import fitz  # PyMuPDF
import re

# ========== CONFIG ========== #
TWITTER_CHAR_LIMIT = 280  # Updated character limit
# Generic hashtags, the AI will add specific ones.
# PREDEFINED_HASHTAGS = [
#     "#Tech", "#Event", "#Learning", "#Workshop", "#Community", "#SkillDevelopment"
# ]

# ========== FILE PARSER (SUPPORTS DOCX & PDF) ========== #
def parse_uploaded_file(uploaded_file):
    """
    Extract text from an uploaded file, supporting both DOCX and PDF formats.
    """
    try:
        file_extension = uploaded_file.name.split('.')[-1].lower()
        
        if file_extension == 'docx':
            doc = Document(io.BytesIO(uploaded_file.read()))
            full_text = [p.text for p in doc.paragraphs if p.text.strip()]
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            full_text.append(cell.text.strip())
            content = '\n'.join(full_text)

        elif file_extension == 'pdf':
            pdf_document = fitz.open(stream=uploaded_file.read(), filetype="pdf")
            full_text = [page.get_text() for page in pdf_document]
            content = '\n'.join(full_text)
            
        else:
            return f"Error: Unsupported file type '{file_extension}'. Please upload a DOCX or PDF file."

        if not content.strip():
            return "Error: Could not extract any text from the document."
            
        return content
        
    except Exception as e:
        return f"Error parsing file: {e}. The file might be corrupted."

# ========== PROMPT BUILDERS (UPGRADED) ========== #

def build_linkedin_post_event_prompt(data, report_content):
    """
    Builds a sophisticated prompt that instructs the AI to extract multiple specific sections.
    """
    example_style_guide = """
    **Example of the desired format and tone (do not copy the content, only the style):**
    
    🌟 Event Name – A Huge Success! 🌟

    We’re thrilled to share that our session "Event Name" saw an amazing turnout with [Number] enthusiastic attendees! 🎉

    💡 **What we covered:**
    - Point 1 (e.g., Fundamentals of CP)
    - Point 2 (e.g., Intermediate techniques)
    - Point 3 (e.g., Advanced tools)
    - (Up to 8 points total)

    📌 **Highlights:**
    - A memorable moment from the event.
    - A unique feature of the session.

    🙌 **What attendees are saying:**
    - "A summary of positive feedback."
    - "Another key positive sentiment."
    
    A big thank you to all the speakers - [Name 1, Name 2, Name 3] for sharing their valuable insights. 
    Thank you to the organizing team at [Organization Name] and the participants for making this session a memorable one.
    Let’s keep learning, sharing, and growing together! 🚀

    [Signature Block]

    #[EventHashtag] #[TopicHashtag] #[CommunityHashtag] #[More Relevant Hashtags]
    """

    prompt = f"""
You are an expert social media manager for a student-run tech club. Your task is to create an engaging, professional LinkedIn post by analyzing an event report.

**Your Goal:**
Read the **Full Event Report** provided below and autonomously write a complete social media post that is exciting, professional, and structured.

**Instructions:**
1.  **Analyze the Report**: From the report, you MUST identify and extract information for these three sections:
    * **What we covered**: A bulleted list of the main topics. **Limit this to 6-8 key points.** Use a 💡 emoji for the heading.
    * **Highlights**: A short, bulleted list of **2-3 memorable moments** or unique features (e.g., "Personal advice segments," "Real contest experience"). Use a 📌 emoji for the heading.
    * **Positive Feedback**: A summary of positive feedback from the report, presented as **2-3 short, impactful points.** Use a 🙌 emoji for the heading.
2.  **Write the Post**:
    * **Headline**: Start with a powerful, emoji-filled headline.
    * **Opening**: Write an enthusiastic opening paragraph announcing the event's success, mentioning the event title and attendance.
    * **Body**: Include the "What we covered," "Highlights," and "Positive Feedback" sections you created.
    * **Gratitude**: Write a warm "Thank You" paragraph using the names from the **"People to Thank"** section.
    * **Closing & Signature**: End with a motivational closing line and add the **Signature Block**.
3.  **Hashtags**: Include the predefined hashtags AND generate 4-5 additional, relevant hashtags.
4.  **Follow the Style Guide**: Use the provided example as a strict guide for the **style, tone, and structure**.

---
**Information to Incorporate:**
* **People to Thank (Speakers, Organizers, etc.)**: {data['people_to_thank']}
* **Signature Block (Heads to Tag)**: {data['signature']}

---
**Full Event Report to Analyze:**
{report_content}
---
{example_style_guide}
---

Now, generate the complete and final LinkedIn post draft.
Generated post will directly be posted so ensure it is polished, professional, and engaging and in the event organizers prespective.
"""
    return prompt

def build_instagram_whatsapp_prompt(platform, data, report_content):
    tone_instruction = "fun, witty, and highly visual, using plenty of relevant emojis (✨, 🚀, 📸, 🙌)." if platform == "Instagram" else "friendly, celebratory, and clear."

    prompt = f"""
You are a social media manager for a student tech club. Create a {platform} post by analyzing the event report below.

**Instructions:**
1.  **Analyze the Report**: Read the report to find the event title, attendance, and the most exciting highlights.
2.  **Set the Tone**: The tone must be {tone_instruction}
3.  **Write the Post**:
    * Start with a catchy, emoji-filled title.
    * Write a short, energetic paragraph about the event's success.
    * Create a bulleted list of 3-5 key highlights or topics covered.
    * If the report mentions positive feedback, add one or two quotes or summarized points.
    * Give a big "Thank You" to the people mentioned in the **"People to Thank"** section.
    * End with a short, motivational closing line.
4.  **Hashtags (for Instagram)**: If the platform is Instagram, include a mix of community and topic-specific hashtags.

---
**Information to Incorporate:**
* **People to Thank**: {data['people_to_thank']}
* **Signature/Tag**: {data['signature']}
---
**Full Event Report to Analyze:**
{report_content}
---

Now, generate the complete {platform} post draft.
Generated post will directly be posted so ensure it is polished, professional, and engaging and in the event organizers prespective.
"""
    return prompt

def build_twitter_prompt(data, report_content):
    prompt = f"""
Generate a post-event tweet (X post) under {TWITTER_CHAR_LIMIT} characters total.

Event details:
Title: 
Report Summary: 
Date: 
Organizing Team: 
Speakers: 
Attendance:

Tweet MUST include:
- A short, celebratory wrap-up message.
- Mention of event title and date.
- Thank-you note to the organizing team and speakers.
- ~25-character Instagram post link included in this format: (📸 See more: https://shorturl.at/XXXX)
- Hashtags: 
- The tone should be energetic, club-friendly, and tweet-length conscious.
- The total character count (including hashtags and link) MUST be <= {TWITTER_CHAR_LIMIT}.
- Place the Instagram link toward the end of the tweet.
- Output only the tweet text. No extra notes or explanations.
**Full Event Report to Analyze:**
{report_content}
"""
    return prompt

# ========== GROQ API & REFINEMENT PROCESS ========== #
def call_groq_api_with_refinement(api_key, initial_prompt, platform):
    """
    Generates a post in a two-step process: first a draft, then a refinement.
    """
    try:
        client = Groq(api_key=api_key)
        
        # Step 1: Generate the initial draft
        draft_completion = client.chat.completions.create(
            model="llama3-70b-8192",
            messages=[{"role": "user", "content": initial_prompt}],
            temperature=0.7,
            top_p=0.9
        )
        draft_post = draft_completion.choices[0].message.content.strip()

        if platform == "Twitter":
            return draft_post[:TWITTER_CHAR_LIMIT]

        # Step 2: Refine the draft for other platforms
        refinement_prompt = f"""
You are an expert social media copy editor. Your task is to polish the following draft for a {platform} post.

**Instructions:**
1.  **Improve Flow**: Ensure the text flows naturally and is easy to read.
2.  **Enhance Emojis**: Add engaging and relevant emojis (like 🌟, 🎉, 🚀, ✅, 💡, 📌, 🙌) to make the post visually appealing and professional.
3.  **Check Professionalism**: Ensure the tone is appropriate for {platform}.
4.  **Formatting**: Ensure lists are correctly formatted and use bolding for emphasis on key phrases and headings.
5.  **Do NOT Change Core Information**: Only improve the presentation of the draft.

**Draft to Refine:**
---
{draft_post}
---

Return only the final, polished post.
Generated post will directly be posted so ensure it is polished, professional, and engaging and in the event organizers prespective.
"""
        
        refined_completion = client.chat.completions.create(
            model="llama3-70b-8192",
            messages=[{"role": "user", "content": refinement_prompt}],
            temperature=0.5,
        )
        return refined_completion.choices[0].message.content.strip()

    except Exception as e:
        st.error(f"An error occurred with the Groq API: {e}")
        return "Error: Could not generate content."

# ========== STREAMLIT UI ========== #
st.set_page_config(page_title="Post-Event Content Generator", layout="wide", page_icon="🎉")
st.title("🎉 AI Post-Event Content Generator")
st.markdown("Generate high-quality social media content by uploading your event report. The AI will read the report and write the posts for you.")

with st.sidebar:
    st.header("⚙️ Configuration")
    api_key = st.text_input("Groq API Key", type="password")
    st.markdown("---")
    st.header("📄 Upload Report")
    uploaded_file = st.file_uploader("Upload Report (DOCX or PDF)", type=['docx', 'pdf'])

if not uploaded_file:
    st.info("Please upload your event report in the sidebar to get started.")
    st.stop()

with st.spinner("📖 Reading and analyzing document..."):
    report_content = parse_uploaded_file(uploaded_file)

if "Error:" in report_content:
    st.error(report_content)
    st.stop()
else:
    st.sidebar.success("✅ Report uploaded successfully!")
    with st.sidebar.expander("📋 Preview Report Content"):
        st.text_area("Document Text", report_content, height=200, disabled=True, key="report_preview")

st.header("✍️ Add Final Details")
st.caption("Provide the names for acknowledgements. The AI will handle the rest.")

with st.form("final_details_form"):
    people_to_thank = st.text_area(
        "Whom to Thank*",
        placeholder="List the names of speakers, organizers, etc., to be thanked in the post.",
        height=100
    )
    signature = st.text_area(
        "Signature / Heads to Tag*",
        placeholder="Enter the signature block, e.g., 'Dr. Geetanjali Kale, ACM India'",
        height=80
    )
    
    submit = st.form_submit_button("✨ Generate All Social Media Content", use_container_width=True)

if submit:
    if not api_key:
        st.error("🔑 Please provide your Groq API Key in the sidebar.")
    elif not all([people_to_thank, signature]):
        st.warning("⚠️ Please fill out both the 'Whom to Thank' and 'Signature' fields.")
    else:
        data = {"people_to_thank": people_to_thank, "signature": signature}
        results = {}
        
        with st.spinner("🤖 Generating and refining professional social media content..."):
            results["LinkedIn"] = call_groq_api_with_refinement(api_key, build_linkedin_post_event_prompt(data, report_content), "LinkedIn")
            results["Instagram"] = call_groq_api_with_refinement(api_key, build_instagram_whatsapp_prompt("Instagram", data, report_content), "Instagram")
            results["WhatsApp"] = call_groq_api_with_refinement(api_key, build_instagram_whatsapp_prompt("WhatsApp", data, report_content), "WhatsApp")
            results["Twitter"] = call_groq_api_with_refinement(api_key, build_twitter_prompt(data, report_content), "Twitter")

        if any(results.values()):
            st.success("🎉 Content generated successfully!")
            tab1, tab2, tab3, tab4 = st.tabs(["📱 LinkedIn", "📸 Instagram", "💬 WhatsApp", "🐦 Twitter (X)"])
            
            with tab1:
                st.subheader("LinkedIn Post")
                st.text_area("LinkedIn Content", value=results.get("LinkedIn", ""), height=500, key="linkedin_post")
            with tab2:
                st.subheader("Instagram Caption")
                st.text_area("Instagram Content", value=results.get("Instagram", ""), height=400, key="instagram_post")
            with tab3:
                st.subheader("WhatsApp Message")
                st.text_area("WhatsApp Content", value=results.get("WhatsApp", ""), height=400, key="whatsapp_post")
            with tab4:
                st.subheader("Twitter (X) Post")
                tweet = results.get("Twitter", "")
                st.text_area("Twitter Content", value=tweet, height=200, key="twitter_post")
                st.caption(f"Character count: {len(tweet)} / {TWITTER_CHAR_LIMIT}")
