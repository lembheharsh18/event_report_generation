from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from io import BytesIO
from charts import create_attendance_chart, create_sentiment_chart, create_word_frequency_chart
import re

# Helper functions for document creation
def add_heading(document, text, level=1, indent=None):
    if text.strip():
        heading = document.add_heading(level=level)
        run = heading.add_run(text.strip())
        run.font.size = Pt(16 if level == 1 else 14 if level == 2 else 12)
        run.font.color.rgb = RGBColor(0, 0, 0)
        run.bold = True
        if indent is not None:
            heading.paragraph_format.left_indent = indent

def add_paragraph(document, text, bold=False, italic=False, font_size=12, indent=None):
    if text.strip():
        p = document.add_paragraph()
        run = p.add_run(text.strip())
        run.font.size = Pt(font_size)
        run.bold = bold
        run.italic = italic
        if indent is not None:
            p.paragraph_format.left_indent = indent

def add_bullet_list(document, items, heading=None, indent=Inches(0.5), bullet_indent=Inches(1)):
    if heading:
        p = document.add_paragraph()
        # Parse Markdown bold in heading
        parts = re.split(r'(\*\*.*?\*\*)', heading)
        for part in parts:
            run = p.add_run(part[2:-2]) if part.startswith('**') and part.endswith('**') and len(part) > 4 else p.add_run(part)
            run.bold = part.startswith('**') and part.endswith('**') and len(part) > 4
            run.font.size = Pt(12)
        if indent is not None:
            p.paragraph_format.left_indent = indent
    for item in items:
        if item.strip():
            # Remove leading number-dot pattern (e.g., '1. ', '2. ', etc.)
            cleaned_item = re.sub(r'^\d+\.\s*', '', item.strip())
            # If the item contains a colon, bold the part before the colon
            if ':' in cleaned_item:
                before, after = cleaned_item.split(':', 1)
                # Remove '**' immediately after the colon in 'after'
                after = re.sub(r'^\*\*', '', after)
                cleaned_item = f"**{before.strip()}:**{after}"
            p = document.add_paragraph(style='List Bullet')
            # Parse Markdown bold in bullet item
            parts = re.split(r'(\*\*.*?\*\*)', cleaned_item)
            for part in parts:
                run = p.add_run(part[2:-2]) if part.startswith('**') and part.endswith('**') and len(part) > 4 else p.add_run(part)
                run.bold = part.startswith('**') and part.endswith('**') and len(part) > 4
                run.font.size = Pt(11)
            if bullet_indent is not None:
                p.paragraph_format.left_indent = bullet_indent

def add_bullet_list_with_headings(document, text, indent=Inches(0.5), bullet_indent=Inches(1)):
    """
    Parses text with headings and bullet points, adds bold headings and bullet lists to the document.
    If no heading is detected, adds the bullet list with a default heading '**Here are suggestions:**'.
    """
    lines = text.splitlines()
    current_heading = None
    bullets = []
    any_heading = False
    for line in lines:
        heading_match = re.match(r"^(\*\*.*\*\*|#|\*|\-|\d+\.|[A-Za-z ]+:)", line.strip())
        if heading_match and (line.strip().endswith(":") or (line.strip().startswith("**") and line.strip().endswith("**"))):
            # Flush previous bullets
            if bullets:
                # If no heading so far, use a default heading
                add_bullet_list(document, bullets, heading=current_heading or ("**Here are suggestions:**" if not any_heading else None), indent=indent, bullet_indent=bullet_indent)
                bullets = []
            # Clean heading
            heading = line.strip().rstrip(":").strip("*# ")
            # Ensure heading is bolded in Markdown
            if not (heading.startswith("**") and heading.endswith("**")):
                heading = f"**{heading}**"
            current_heading = heading
            any_heading = True
        elif line.strip().startswith("-") or line.strip().startswith("*"):
            # Remove leading number-dot pattern from bullet
            bullet = line.strip().lstrip("-* ")
            bullet = re.sub(r'^\d+\.\s*', '', bullet)
            bullets.append(bullet)
        elif line.strip():
            bullet = line.strip()
            bullet = re.sub(r'^\d+\.\s*', '', bullet)
            bullets.append(bullet)
    # Flush last
    if bullets:
        add_bullet_list(document, bullets, heading=current_heading or ("**Here are suggestions:**" if not any_heading else None), indent=indent, bullet_indent=bullet_indent)

def add_chart(document, chart_data, caption, width=5):
    if chart_data:
        chart_bytes = chart_data.read()
        document.add_picture(BytesIO(chart_bytes), width=Inches(width))
        last_paragraph = document.paragraphs[-1] 
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_paragraph(document, caption, italic=True, font_size=10)

def add_horizontal_line(document):
    # Add an empty paragraph
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Use oxml to add a true bottom border (horizontal line)
    p_elm = p._p
    pPr = p_elm.get_or_add_pPr()
    pbdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '12')  # thickness
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '000000')
    pbdr.append(bottom)
    pPr.append(pbdr)

def add_markdown_paragraph(document, text, font_size=12, indent=None):
    p = document.add_paragraph()
    if indent is not None:
        p.paragraph_format.left_indent = indent
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**') and len(part) > 4:
            run = p.add_run(part[2:-2])
            run.bold = True
        else:
            run = p.add_run(part)
        run.font.size = Pt(font_size)

def create_docx_report(report_data, analysis, analyzer):
    doc = Document()
    def safe_professionalize(text):
        """Helper to sanitize Groq output and remove chatty filler content."""
        result = analyzer.professionalize_text(text).strip()
        bad_phrases = [
            "here is a rewritten version", "i apologize", "please provide", "i'll be happy",
            "the original text", "if you meant", "feel free", "it appears", "i assumed", 
            "there is no content", "let me know", "make any adjustments"
        ]
        result_lower = result.lower()
        if any(phrase in result_lower for phrase in bad_phrases):
            return text.strip()  # fallback to raw original input
        return result

    # Professionalize event details
    event_name = report_data.get('event_name', '')
    event_venue = report_data.get('event_venue', '')

    # Add title
    title = doc.add_paragraph()
    title_run = title.add_run(f"Event Report: {event_name}")
    title_run.font.size = Pt(20)
    title_run.bold = True
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER


    # Section 1: Event Details
    add_heading(doc, "📋 Event Details", level=1)
    if event_name:
        add_markdown_paragraph(doc, f"**Event Name:** {event_name}", font_size=11, indent=Inches(0.5))
    if report_data.get('event_date'):
        add_markdown_paragraph(doc, f"**Date:** {report_data['event_date']}", font_size=11, indent=Inches(0.5))
    if event_venue:
        add_markdown_paragraph(doc, f"**Venue:** {event_venue}", font_size=11, indent=Inches(0.5))
    if report_data.get('event_time'):
        add_markdown_paragraph(doc, f"**Time:** {report_data['event_time']}", font_size=11, indent=Inches(0.5))

    add_markdown_paragraph(doc, f"**Total Attendees:** {report_data.get('total_attendees', 'N/A')}", font_size=11, indent=Inches(0.5))
    add_markdown_paragraph(
        doc,
        f"**Attendee Breakdown:** Students ({report_data.get('students', '0')}), Faculty ({report_data.get('faculty', '0')}), Guests ({report_data.get('guests', '0')})",
        font_size=11, indent=Inches(0.5)
    )

    # ** CORRECTION STARTS HERE **
    # Attendance chart
    # Helper function to safely parse the first integer from a string like '50+' or 'Around 100'
    def _parse_num(s):
        # Find all sequences of digits in the string
        nums = re.findall(r'\d+', str(s))
        # Return the first number found, or 0 if no numbers are found
        return int(nums[0]) if nums else 0

    students_num = _parse_num(report_data.get('students', '0'))
    faculty_num = _parse_num(report_data.get('faculty', '0'))
    guests_num = _parse_num(report_data.get('guests', '0'))
    
    # Check the sum of parsed numbers to decide if a chart should be created
    if students_num + faculty_num + guests_num > 0:
        attendance_chart = create_attendance_chart(
            report_data.get('students', '0'),
            report_data.get('faculty', '0'),
            report_data.get('guests', '0')
        )
        add_chart(doc, attendance_chart, "Figure 1: Attendee Breakdown")
    # ** CORRECTION ENDS HERE **
    add_horizontal_line(doc)

    # Section 2: Topics Covered
    add_heading(doc, "📚 Topics Covered", level=1)
    if report_data.get('topics'):
        prof_topics = [safe_professionalize(t) for t in report_data['topics']]
        prof_topics = [t for t in prof_topics if t]
        if prof_topics:
            for t in prof_topics:
                add_markdown_paragraph(doc, f"**{t}**", font_size=11, indent=Inches(0.5))
        else:
            add_paragraph(doc, "No valid topics recorded", italic=True, indent=Inches(0.5))
    else:
        add_paragraph(doc, "No topics recorded", italic=True, indent=Inches(0.75))
    add_horizontal_line(doc)

    # Section 3: Feedback Analysis
    add_heading(doc, "📊 Feedback Analysis", level=1)
    sentiment_data = analysis.get("sentiment_analysis", {})
    if sentiment_data:
        add_heading(doc, "😃 Sentiment Analysis", level=2, indent=Inches(0.5))
        add_markdown_paragraph(doc, f"**Total Responses:** {analysis.get('total_responses', 0)}", font_size=11, indent=Inches(0.75))
        add_markdown_paragraph(doc, f"**Positive:** {sentiment_data.get('positive', 0)} ({sentiment_data.get('percentages', {}).get('positive', 0):.1f}%)", font_size=11, indent=Inches(0.75))
        add_markdown_paragraph(doc, f"**Negative:** {sentiment_data.get('negative', 0)} ({sentiment_data.get('percentages', {}).get('negative', 0):.1f}%)", font_size=11, indent=Inches(0.75))
        add_markdown_paragraph(doc, f"**Neutral:** {sentiment_data.get('neutral', 0)} ({sentiment_data.get('percentages', {}).get('neutral', 0):.1f}%)", font_size=11, indent=Inches(0.75))
        add_markdown_paragraph(doc, f"**Overall Sentiment Score:** {sentiment_data.get('overall_score', 0):.2f}", font_size=11, indent=Inches(0.75))
        sentiment_chart = create_sentiment_chart({
            "Positive": sentiment_data.get('positive', 0),
            "Neutral": sentiment_data.get('neutral', 0),
            "Negative": sentiment_data.get('negative', 0)
        })
        add_chart(doc, sentiment_chart, "Figure 2: Feedback Sentiment Distribution")
    

    # Text analysis
    text_analysis = analysis.get("text_analysis", {})
    if text_analysis:
        add_heading(doc, "📝 Text Analysis", level=2, indent=Inches(0.5))
        add_markdown_paragraph(doc, f"**Total Words:** {text_analysis.get('total_words', 0)}", font_size=11, indent=Inches(0.75))
        add_markdown_paragraph(doc, f"**Unique Words:** {text_analysis.get('unique_words', 0)}", font_size=11, indent=Inches(0.75))
        if text_analysis.get('most_common_words'):
            word_chart = create_word_frequency_chart(text_analysis['most_common_words'])
            add_chart(doc, word_chart, "Figure 3: Top 20 Most Common Words", width=6)
    

    add_heading(doc, "🧾 Summary", level=2, indent=Inches(0.5))
    summary = analysis.get("narrative_summary", "No summary available")
    if summary and summary != "No summary available":
        add_markdown_paragraph(doc, summary, font_size=11, indent=Inches(0.75))
    else:
        add_paragraph(doc, summary, font_size=11, indent=Inches(0.75))
    

    add_heading(doc, "💡 Key Takeaways", level=2, indent=Inches(0.5))
    takeaways = analysis.get("key_takeaways", "No takeaways available")
    add_bullet_list_with_headings(doc, takeaways, indent=Inches(0.75), bullet_indent=Inches(1))
    add_horizontal_line(doc)
    # Section 4: Challenges & Solutions
    if report_data.get('challenges') or report_data.get('solutions'):
        add_heading(doc, "⚠️ Challenges Faced and Solutions", level=1)
        # Ensure challenges and solutions are of the same length for zipping
        challenges = report_data.get('challenges', [])
        solutions = report_data.get('solutions', [])
        max_len = max(len(challenges), len(solutions))
        challenges.extend([''] * (max_len - len(challenges)))
        solutions.extend([''] * (max_len - len(solutions)))

        for i, (ch, sol) in enumerate(zip(challenges, solutions)):
            prof_ch = safe_professionalize(ch)
            prof_sol = safe_professionalize(sol)
            if prof_ch or prof_sol:
                add_heading(doc, f"⚠️ Challenge {i+1}", level=2, indent=Inches(0.5))
                if prof_ch:
                    add_paragraph(doc, prof_ch, font_size=11, indent=Inches(0.75))
                add_heading(doc, f"🛠️ Solution {i+1}", level=2, indent=Inches(0.5))
                if prof_sol:
                    add_paragraph(doc, prof_sol, font_size=11, indent=Inches(0.75))
        add_horizontal_line(doc)

    # Section 5: Suggestions
    add_heading(doc, "✨ Suggestions for Future Events", level=1)
    suggestions = analysis.get("suggestions", "No suggestions available")
    add_bullet_list_with_headings(doc, suggestions, indent=Inches(0.75))
    add_horizontal_line(doc)

    # Section 6: Special Mentions
    if report_data.get('special_mentions') or report_data.get('relevant_links'):
        add_heading(doc, "🌟 Special Mentions & Links", level=1)
        
        if report_data.get('special_mentions'):
            add_bullet_list(doc, report_data['special_mentions'], heading="Acknowledgements", indent=Inches(0.75))
        
        if report_data.get('relevant_links'):
            add_bullet_list(doc, report_data['relevant_links'], heading="Relevant Resources", indent=Inches(0.75))
        add_horizontal_line(doc)

    # Footer
    prepared_by = report_data.get('prepared_by')
    if prepared_by:
        add_heading(doc, "✍️ Report Prepared By", level=1)
        add_paragraph(doc, f"Name: {prepared_by}", font_size=12, indent=Inches(0.75))
    if report_data.get('report_date'):
        add_paragraph(doc, f"Date: {report_data['report_date']}", font_size=12, indent=Inches(0.75))

    # Save DOCX
    doc_bytes = BytesIO()
    doc.save(doc_bytes)
    doc_bytes.seek(0)
    return doc_bytes
